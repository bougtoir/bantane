#!/usr/bin/env python3
"""管理者向けマニュアル（難読化ビルド・ライセンス管理）の docx を生成する。"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime as dt


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, font_size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)

    doc.add_paragraph()  # spacing
    return table


def main():
    doc = Document()

    # ---- Title ----
    title = doc.add_heading('ばんたね病院 シフト最適化アプリ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('管理者マニュアル — 難読化ビルド・ライセンス管理', level=0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)

    add_para(doc, f'最終更新: {dt.date.today().strftime("%Y年%m月%d日")}',
             italic=True, font_size=10)
    doc.add_paragraph()

    # ---- 目次 ----
    add_heading(doc, '目次', level=1)
    toc_items = [
        '1. 概要',
        '2. 前提条件',
        '3. Nuitka による難読化ビルド',
        '4. ライセンス管理',
        '5. 配布手順',
        '6. ライセンス更新・再発行',
        '7. トラブルシューティング',
        '8. セキュリティに関する注意事項',
    ]
    for item in toc_items:
        add_para(doc, item)
    doc.add_page_break()

    # ==== 1. 概要 ====
    add_heading(doc, '1. 概要', level=1)
    add_para(doc, (
        '本マニュアルは、シフト最適化アプリの管理者（IT担当者）を対象に、'
        'ソースコードの難読化ビルドとライセンス管理の手順を説明します。'
    ))
    add_para(doc, '本アプリでは以下の2つの保護機構を導入しています：')
    add_bullet(doc, '難読化ビルド（Nuitka）: Python ソースコードを C 言語にコンパイルし、'
                     'ネイティブ実行ファイル（.exe）を生成。逆コンパイルを極めて困難にします。')
    add_bullet(doc, 'ライセンス管理: AES-128 暗号化 + ハードウェアバインドにより、'
                     '不正コピーと無断使用を防止します。')

    # ==== 2. 前提条件 ====
    add_heading(doc, '2. 前提条件', level=1)
    add_heading(doc, '2.1 ビルド環境', level=2)
    add_table(doc,
              ['項目', '要件'],
              [
                  ['OS', 'Windows 10/11 (64-bit)'],
                  ['Python', '3.10 以上'],
                  ['C コンパイラ', 'MinGW64 または Visual Studio Build Tools'],
                  ['ディスク空き容量', '約 5 GB（Nuitka コンパイル用）'],
                  ['メモリ', '8 GB 以上推奨'],
              ])

    add_heading(doc, '2.2 必要なパッケージ', level=2)
    add_para(doc, 'requirements.txt に加えて、ビルドには以下が必要です：')
    add_bullet(doc, 'nuitka — Python → C コンパイラ')
    add_bullet(doc, 'ordered-set — Nuitka のコンパイル高速化（推奨）')
    add_bullet(doc, 'cryptography — ライセンス暗号化（Fernet / AES-128-CBC）')
    add_para(doc, 'これらは build_nuitka.bat が自動でインストールします。')

    add_heading(doc, '2.3 C コンパイラのインストール', level=2)
    add_para(doc, 'Nuitka は C コンパイラを必要とします。以下のいずれかをインストールしてください：')
    add_para(doc, '方法 A: MinGW64（推奨 — 軽量）', bold=True)
    add_code_block(doc, 'python -m nuitka --mingw64  # 初回実行時に自動ダウンロード')
    add_para(doc, '方法 B: Visual Studio Build Tools', bold=True)
    add_bullet(doc, 'https://visualstudio.microsoft.com/visual-cpp-build-tools/ からダウンロード')
    add_bullet(doc, '「C++ によるデスクトップ開発」ワークロードを選択してインストール')

    doc.add_page_break()

    # ==== 3. Nuitka 難読化ビルド ====
    add_heading(doc, '3. Nuitka による難読化ビルド', level=1)

    add_heading(doc, '3.1 ビルド手順', level=2)
    add_para(doc, '1. プロジェクトフォルダで build_nuitka.bat をダブルクリックします。')
    add_para(doc, '2. 自動的に以下の処理が実行されます：')
    add_bullet(doc, '仮想環境の作成（初回のみ）')
    add_bullet(doc, '依存パッケージのインストール')
    add_bullet(doc, 'Nuitka のインストール')
    add_bullet(doc, 'Python → C → ネイティブ exe のコンパイル')
    add_para(doc, '3. ビルド完了後、以下のファイルが生成されます：')
    add_code_block(doc, 'dist_nuitka\\BantaneShiftOptimizer.exe')
    add_para(doc, '4. release フォルダに自動的にコピーされます。')

    add_heading(doc, '3.2 ビルド所要時間', level=2)
    add_table(doc,
              ['条件', '目安時間'],
              [
                  ['初回ビルド', '10 〜 20 分'],
                  ['2 回目以降（キャッシュあり）', '3 〜 8 分'],
                  ['コード変更なし（再ビルド）', '1 〜 3 分'],
              ])

    add_heading(doc, '3.3 PyInstaller ビルド（従来方式）', level=2)
    add_para(doc, (
        '従来の build.bat（PyInstaller）も引き続き利用可能です。'
        'ただし PyInstaller は .pyc をバンドルするだけのため、'
        '逆コンパイルツールで容易にソースコードを復元できます。'
        'セキュリティが必要な配布には Nuitka ビルドを推奨します。'
    ))

    add_heading(doc, '3.4 難読化の比較', level=2)
    add_table(doc,
              ['方式', '逆コンパイル耐性', 'ビルド速度', '実行速度'],
              [
                  ['PyInstaller (build.bat)', '低（.pyc 復元可能）', '速い（1-2分）', '通常'],
                  ['Nuitka (build_nuitka.bat)', '高（C コンパイル）', '遅い（10-20分）', 'やや高速'],
              ])

    doc.add_page_break()

    # ==== 4. ライセンス管理 ====
    add_heading(doc, '4. ライセンス管理', level=1)

    add_heading(doc, '4.1 ライセンスの仕組み', level=2)
    add_para(doc, '本アプリのライセンスは以下の3つの要素で保護されています：')
    add_bullet(doc, 'ユーザーID / パスワード認証 — アプリ起動時に入力')
    add_bullet(doc, 'AES-128 暗号化（Fernet）— ライセンスファイルの暗号化')
    add_bullet(doc, 'ハードウェアバインド — マシン固有の情報（ホスト名 + MAC アドレス + '
                     'ボリュームシリアル番号）を SHA-256 ハッシュ化し、ライセンスに紐付け')

    add_para(doc, (
        'ライセンスファイル（.license）を別の PC にコピーしても、'
        'ハードウェアフィンガープリントが一致しないため認証が失敗します。'
    ))

    add_heading(doc, '4.2 ライセンス発行（generate_license.py）', level=2)
    add_para(doc, '管理者は generate_license.py を使ってライセンスファイルを発行します。')

    add_para(doc, '方法 A: 対話モード', bold=True)
    add_code_block(doc, (
        'python generate_license.py\n'
        '# → ユーザーID、パスワード、マシンID を順に入力'
    ))

    add_para(doc, '方法 B: コマンドライン', bold=True)
    add_code_block(doc, (
        'python generate_license.py \\\n'
        '    --user-id tanaka \\\n'
        '    --password shift2026 \\\n'
        '    --days 365 \\\n'
        '    --machine-id <対象PCのマシンID>'
    ))

    add_para(doc, '方法 C: 現在の PC にライセンス発行', bold=True)
    add_code_block(doc, (
        'python generate_license.py \\\n'
        '    --user-id tanaka \\\n'
        '    --password shift2026 \\\n'
        '    --days 365'
    ))

    add_heading(doc, '4.3 マシンID の確認方法', level=2)
    add_para(doc, '管理者側:')
    add_code_block(doc, 'python generate_license.py --show-machine-id')
    add_para(doc, 'エンドユーザー側:')
    add_para(doc, (
        'アプリの「ライセンス認証」画面に「マシンID」が表示されます。'
        '「コピー」ボタンを押してクリップボードにコピーし、管理者にメールやチャットで送信してください。'
    ))

    add_heading(doc, '4.4 コマンドラインオプション一覧', level=2)
    add_table(doc,
              ['オプション', '説明', 'デフォルト'],
              [
                  ['--user-id', 'ユーザーID', '（対話入力）'],
                  ['--password', 'パスワード', '（対話入力）'],
                  ['--days', '有効日数', '365'],
                  ['--machine-id', '対象マシンのフィンガープリント', '（現在のPC）'],
                  ['--output', '出力先ファイルパス', '.license'],
                  ['--show-machine-id', '現在のPCのマシンIDを表示', '—'],
              ])

    doc.add_page_break()

    # ==== 5. 配布手順 ====
    add_heading(doc, '5. 配布手順', level=1)

    add_heading(doc, '5.1 初回配布', level=2)
    add_para(doc, '以下の手順でエンドユーザーにアプリを配布します：')
    add_para(doc, '手順 1: ビルド', bold=True)
    add_bullet(doc, 'build_nuitka.bat を実行して exe をビルド')
    add_para(doc, '手順 2: マシンID の取得', bold=True)
    add_bullet(doc, 'エンドユーザーの PC でマシンID を確認')
    add_bullet(doc, '方法 a: generate_license.py --show-machine-id を対象 PC で実行')
    add_bullet(doc, '方法 b: アプリを一度起動してもらい、認証画面のマシンID をコピーしてもらう')
    add_para(doc, '手順 3: ライセンス発行', bold=True)
    add_bullet(doc, 'generate_license.py でライセンスファイルを作成（マシンID を指定）')
    add_para(doc, '手順 4: ファイル一式を配布', bold=True)
    add_para(doc, '以下のファイルをフォルダにまとめてエンドユーザーに渡します：')
    add_table(doc,
              ['ファイル', '説明'],
              [
                  ['BantaneShiftOptimizer.exe', 'アプリ本体（Nuitka ビルド済み）'],
                  ['.license', 'ライセンスファイル（マシンバインド済み）'],
                  ['files/setting_YYYY_MM.xlsx', '業務設定ファイル'],
              ])

    add_heading(doc, '5.2 フォルダ構成例', level=2)
    add_code_block(doc, (
        '配布フォルダ/\n'
        '├── BantaneShiftOptimizer.exe\n'
        '├── .license\n'
        '└── files/\n'
        '    └── setting_2026_07.xlsx'
    ))

    # ==== 6. ライセンス更新 ====
    add_heading(doc, '6. ライセンス更新・再発行', level=1)

    add_heading(doc, '6.1 有効期限の更新', level=2)
    add_para(doc, (
        'ライセンスの有効期限が切れた場合、新しいライセンスファイルを発行して '
        '古い .license ファイルを上書きしてください。'
    ))
    add_code_block(doc, (
        'python generate_license.py \\\n'
        '    --user-id tanaka \\\n'
        '    --password shift2026 \\\n'
        '    --days 365 \\\n'
        '    --machine-id <マシンID> \\\n'
        '    --output .license'
    ))

    add_heading(doc, '6.2 PC 入れ替え時', level=2)
    add_para(doc, (
        'エンドユーザーの PC が変わった場合は、新しい PC のマシンID を取得し、'
        'そのマシンID で新しいライセンスを発行してください。'
        '古い PC のライセンスは自動的に無効になります（ハードウェアバインドのため）。'
    ))

    add_heading(doc, '6.3 パスワード変更', level=2)
    add_para(doc, (
        'パスワードを変更する場合は、新しいパスワードでライセンスを再発行し、'
        'エンドユーザーに新しいパスワードと .license ファイルの両方を渡してください。'
    ))

    doc.add_page_break()

    # ==== 7. トラブルシューティング ====
    add_heading(doc, '7. トラブルシューティング', level=1)

    add_table(doc,
              ['症状', '原因', '対処法'],
              [
                  [
                      '「このPCのライセンスではありません」',
                      'ライセンスが別の PC 用に発行されている',
                      'エンドユーザーの PC のマシンID で再発行',
                  ],
                  [
                      '「ライセンスの有効期限が切れています」',
                      '有効期間が過ぎている',
                      '--days オプションで新しい有効日数を指定して再発行',
                  ],
                  [
                      '「ライセンスファイルが見つかりません」',
                      '.license ファイルが exe と同じフォルダにない',
                      '.license ファイルを exe と同じフォルダに配置',
                  ],
                  [
                      '「ユーザーIDが正しくありません」',
                      'ライセンス発行時と異なる ID を入力',
                      '正しいユーザーID を確認して再入力',
                  ],
                  [
                      '「パスワードが正しくありません」',
                      'ライセンス発行時と異なるパスワード',
                      '正しいパスワードを確認。不明な場合は再発行',
                  ],
                  [
                      'Nuitka ビルドが失敗する',
                      'C コンパイラが未インストール',
                      'MinGW64 または Visual Studio Build Tools をインストール',
                  ],
                  [
                      'ビルドは成功するが exe が起動しない',
                      '依存ファイルの欠落',
                      '--include-package や --include-module オプションを追加',
                  ],
              ])

    # ==== 8. セキュリティ ====
    add_heading(doc, '8. セキュリティに関する注意事項', level=1)

    add_bullet(doc, 'ライセンス発行ツール（generate_license.py）および license_manager.py は'
                     '管理者のみがアクセスできる場所に保管してください。'
                     'エンドユーザーには配布しないでください。')
    add_bullet(doc, 'パスワードは十分な強度のものを設定してください（8文字以上、英数字混合推奨）。')
    add_bullet(doc, 'Nuitka でビルドした exe は逆コンパイルが極めて困難ですが、'
                     '完全に不可能ではありません。機密性の高いロジックは別途サーバーサイドで'
                     '処理することも検討してください。')
    add_bullet(doc, 'ライセンスファイルの暗号化鍵（マスターキー）は license_manager.py に'
                     'ハードコードされています。ソースコードの管理には十分注意してください。')

    # ---- Save ----
    output_path = 'bantane_admin_manual.docx'
    doc.save(output_path)
    print(f'マニュアルを生成しました: {output_path}')


if __name__ == '__main__':
    main()
